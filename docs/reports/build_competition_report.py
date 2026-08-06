#!/usr/bin/env python3
"""Build the competition-ready Codebase Memory MCP report."""

from __future__ import annotations

import argparse
import csv
import math
import random
import sys
from collections import defaultdict
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = "#0B2545"
BLUE = "#2E74B5"
DARK_BLUE = "#1F4D78"
TEAL = "#1F7A8C"
GOLD = "#B17A1B"
GREEN = "#2A7D55"
INK = "#1D2733"
MUTED = "#5F6B78"
LIGHT = "#F4F6F9"
LIGHT_BLUE = "#E8EEF5"
LIGHT_TEAL = "#E8F3F4"
LIGHT_GOLD = "#FFF5DF"
BORDER = "#D9E0E7"
WHITE = "#FFFFFF"
RED = "#9B1C1C"
CONTENT_WIDTH_DXA = 9360
FONT_LATIN = "Calibri"
FONT_CJK = "Microsoft YaHei"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value.lstrip("#"))


def set_run_font(
    run,
    *,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    latin: str = FONT_LATIN,
    cjk: str = FONT_CJK,
):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), cjk)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.lstrip("#"))
    shd.set(qn("w:val"), "clear")


def set_cell_border(cell, color: str = BORDER, size: int = 6):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = tc_borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color.lstrip("#"))


def remove_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                element = OxmlElement(f"w:{edge}")
                element.set(qn("w:val"), "nil")
                tc_borders.append(element)


def add_page_field(paragraph):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_LATIN)
    r_fonts.set(qn("w:hAnsi"), FONT_LATIN)
    r_fonts.set(qn("w:eastAsia"), FONT_CJK)
    r_pr.append(r_fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED.lstrip("#"))
    r_pr.append(color)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    r_pr.append(size)
    run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)


def set_image_alt(inline_shape, description: str):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description)


def add_paragraph(
    doc,
    text: str = "",
    *,
    style: str | None = None,
    align=None,
    before: float | None = None,
    after: float | None = None,
    line: float | None = None,
    keep_with_next: bool | None = None,
):
    p = doc.add_paragraph(style=style)
    if text:
        p.add_run(text)
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if line is not None:
        p.paragraph_format.line_spacing = line
    if keep_with_next is not None:
        p.paragraph_format.keep_with_next = keep_with_next
    return p


def add_rich_paragraph(doc, segments, *, style=None, align=None, after=8, line=1.333):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    for text, kwargs in segments:
        run = p.add_run(text)
        set_run_font(run, **kwargs)
    return p


def new_numbering(doc, kind: str = "bullet") -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "280")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "290")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_LATIN)
    r_fonts.set(qn("w:hAnsi"), FONT_LATIN)
    r_fonts.set(qn("w:eastAsia"), FONT_CJK)
    r_pr.append(r_fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list(doc, items, *, kind="bullet"):
    num_id = new_numbering(doc, kind)
    paragraphs = []
    for item in items:
        p = doc.add_paragraph(style="List Paragraph")
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_el)
        p_pr.insert(0, num_pr)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        if isinstance(item, tuple):
            label, detail = item
            r1 = p.add_run(label)
            set_run_font(r1, bold=True, color=NAVY)
            r2 = p.add_run(detail)
            set_run_font(r2, color=INK)
        else:
            run = p.add_run(str(item))
            set_run_font(run, color=INK)
        paragraphs.append(p)
    return paragraphs


def add_callout(doc, text: str, *, label: str | None = None, fill=LIGHT_BLUE, accent=BLUE, geometry=None):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=accent, size=8)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    if label:
        r = p.add_run(f"{label}  ")
        set_run_font(r, bold=True, color=accent, size=10.5)
    r = p.add_run(text)
    set_run_font(r, color=INK, size=10.5)
    geometry(table, [CONTENT_WIDTH_DXA])
    add_paragraph(doc, "", after=2)
    return table


def add_metric_strip(doc, metrics, *, geometry):
    table = doc.add_table(rows=1, cols=len(metrics))
    widths = [CONTENT_WIDTH_DXA // len(metrics)] * len(metrics)
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    for idx, (value, label, note, fill) in enumerate(metrics):
        cell = table.cell(0, idx)
        set_cell_shading(cell, fill)
        set_cell_border(cell, color=BORDER, size=6)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(value)
        set_run_font(r, size=20, bold=True, color=NAVY)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(1)
        set_run_font(p2.add_run(label), size=9.5, bold=True, color=INK)
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(0)
        set_run_font(p3.add_run(note), size=8.5, color=MUTED)
    geometry(table, widths)
    add_paragraph(doc, "", after=2)
    return table


def add_data_table(
    doc,
    headers,
    rows,
    widths,
    *,
    geometry,
    font_size=9.4,
    header_fill=LIGHT,
    numeric_cols=(),
    cell_fills=None,
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, header_fill)
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(str(text)), size=font_size, bold=True, color=NAVY)
    for row_idx, values in enumerate(rows):
        row = table.add_row()
        for col_idx, value in enumerate(values):
            cell = row.cells[col_idx]
            if cell_fills and row_idx < len(cell_fills) and cell_fills[row_idx]:
                set_cell_shading(cell, cell_fills[row_idx])
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            set_run_font(p.add_run(str(value)), size=font_size, color=INK)
    geometry(table, widths)
    add_paragraph(doc, "", after=2)
    return table


def set_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    list_style = styles["List Paragraph"]
    list_style.font.name = FONT_LATIN
    list_style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    list_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    list_style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    list_style.font.size = Pt(11)
    list_style.font.color.rgb = rgb(INK)

    caption = styles["Caption"]
    caption.font.name = FONT_LATIN
    caption._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.keep_with_next = True

    custom = {
        "Cover Kicker": (11, GOLD, True, False, WD_ALIGN_PARAGRAPH.CENTER),
        "Cover Title": (28, NAVY, True, False, WD_ALIGN_PARAGRAPH.CENTER),
        "Cover Subtitle": (14, MUTED, False, False, WD_ALIGN_PARAGRAPH.CENTER),
        "Lead": (13, NAVY, True, False, WD_ALIGN_PARAGRAPH.LEFT),
        "Source Note": (8.5, MUTED, False, False, WD_ALIGN_PARAGRAPH.LEFT),
        "Closing Quote": (15, NAVY, True, False, WD_ALIGN_PARAGRAPH.CENTER),
    }
    for name, (size, color, bold, italic, align) in custom.items():
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = bold
        style.font.italic = italic
        style.paragraph_format.alignment = align
        style.paragraph_format.space_after = Pt(8)


def set_document_geometry(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)


def set_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    left = p.add_run("CODEBASE MEMORY MCP")
    set_run_font(left, size=8.5, bold=True, color=MUTED)
    right = p.add_run("\t参赛项目汇报书")
    set_run_font(right, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    set_run_font(fp.add_run("AI 编码 Agent 的结构化代码记忆层"), size=8.5, color=MUTED)
    set_run_font(fp.add_run("\t"), size=8.5, color=MUTED)
    add_page_field(fp)


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_section_title(
    doc,
    number: str,
    title: str,
    lead: str | None = None,
    *,
    page_break_before: bool = False,
):
    h = doc.add_heading(f"{number}  {title}", level=1)
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.page_break_before = page_break_before
    if lead:
        p = add_paragraph(doc, lead, style="Lead", after=10, line=1.2)
        p.paragraph_format.keep_with_next = False


def font(path: Path, size: int, bold=False):
    filename = "msyhbd.ttc" if bold else "msyh.ttc"
    return ImageFont.truetype(str(path / filename), size=size)


def draw_center(draw, xy, text, fnt, fill, anchor="mm", spacing=6):
    draw.multiline_text(xy, text, font=fnt, fill=fill, anchor=anchor, align="center", spacing=spacing)


def rounded_box(draw, box, fill, outline, radius=24, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def make_cover_graph(path: Path, font_dir: Path):
    w, h = 1600, 580
    image = Image.new("RGB", (w, h), WHITE)
    draw = ImageDraw.Draw(image)
    random.seed(17)
    nodes = []
    center = (800, 290)
    for i in range(22):
        angle = (2 * math.pi * i / 22) + random.uniform(-0.12, 0.12)
        radius = random.choice([150, 225, 300, 390, 480])
        x = int(center[0] + math.cos(angle) * radius)
        y = int(center[1] + math.sin(angle) * radius * 0.48)
        nodes.append((x, y))
    for idx, node in enumerate(nodes):
        draw.line([center, node], fill=BORDER, width=3)
        if idx % 2 == 0:
            other = nodes[(idx + 5) % len(nodes)]
            draw.line([node, other], fill=LIGHT_BLUE, width=2)
    palette = [BLUE, TEAL, GOLD, GREEN]
    for idx, (x, y) in enumerate(nodes):
        radius = 14 if idx % 4 else 20
        draw.ellipse((x - radius, y - radius, x + radius, y + radius), fill=palette[idx % 4], outline=WHITE, width=4)
    rounded_box(draw, (615, 205, 985, 375), NAVY, NAVY, radius=36, width=0)
    draw_center(draw, center, "CODEBASE\nMEMORY", font(font_dir, 42, True), WHITE, spacing=8)
    labels = [
        ((310, 110), "代码实体"),
        ((1295, 140), "调用关系"),
        ((250, 455), "服务接口"),
        ((1335, 440), "变更影响"),
        ((800, 515), "测试链路"),
    ]
    for (x, y), text in labels:
        rounded_box(draw, (x - 85, y - 28, x + 85, y + 28), LIGHT, BORDER, radius=18, width=2)
        draw_center(draw, (x, y), text, font(font_dir, 22, True), NAVY)
    image.save(path, quality=95)


def draw_arrow(draw, start, end, color=BLUE, width=5):
    draw.line([start, end], fill=color, width=width)
    x2, y2 = end
    x1, y1 = start
    angle = math.atan2(y2 - y1, x2 - x1)
    length = 18
    left = (x2 - length * math.cos(angle - 0.45), y2 - length * math.sin(angle - 0.45))
    right = (x2 - length * math.cos(angle + 0.45), y2 - length * math.sin(angle + 0.45))
    draw.polygon([end, left, right], fill=color)


def make_architecture(path: Path, font_dir: Path):
    w, h = 1600, 910
    image = Image.new("RGB", (w, h), WHITE)
    draw = ImageDraw.Draw(image)
    title_font = font(font_dir, 32, True)
    box_font = font(font_dir, 25, True)
    small_font = font(font_dir, 21)
    draw.text((70, 42), "从代码资产到 Agent 可用证据", font=title_font, fill=NAVY)

    stages = [
        (90, 130, 1510, 245, LIGHT, BORDER, "1  输入资产", ["源码与配置", "Git 变更", "IaC 与路由", "运行时 Trace"]),
        (90, 295, 1510, 430, LIGHT_BLUE, BLUE, "2  结构与语义提取", ["文件发现 / 增量检测", "Tree-sitter AST", "Hybrid LSP", "多阶段关系构建"]),
        (90, 480, 1510, 615, LIGHT_TEAL, TEAL, "3  持久化知识图谱", ["节点与限定名", "调用 / 导入 / 类型边", "覆盖度元数据", "SQLite / 共享快照"]),
        (90, 665, 1510, 830, LIGHT_GOLD, GOLD, "4  工具与消费端", ["搜索 / 源码片段", "调用链 / 影响分析", "架构 / Cypher / 跨服务", "MCP / CLI / HTTP / UI"]),
    ]
    for idx, (x1, y1, x2, y2, fill, outline, label, items) in enumerate(stages):
        rounded_box(draw, (x1, y1, x2, y2), fill, outline, radius=26, width=3)
        draw.text((x1 + 28, y1 + 18), label, font=box_font, fill=NAVY)
        card_y1 = y1 + 62
        card_y2 = y2 - 18
        gap = 18
        card_w = (x2 - x1 - 56 - gap * 3) / 4
        for item_idx, item in enumerate(items):
            cx1 = x1 + 28 + item_idx * (card_w + gap)
            cx2 = cx1 + card_w
            rounded_box(draw, (cx1, card_y1, cx2, card_y2), WHITE, BORDER, radius=16, width=2)
            draw_center(draw, ((cx1 + cx2) / 2, (card_y1 + card_y2) / 2), item, small_font, INK)
        if idx < len(stages) - 1:
            draw_arrow(draw, (800, y2 + 5), (800, stages[idx + 1][1] - 5), color=outline, width=5)
    image.save(path, quality=95)


def read_core_results(csv_path: Path):
    with csv_path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def aggregate(rows, dimension):
    result = defaultdict(lambda: {"baseline_tokens": 0, "mcp_tokens": 0, "baseline_ms": 0.0, "mcp_ms": 0.0})
    for row in rows:
        key = row[dimension]
        condition = row["condition"]
        result[key][f"{condition}_tokens"] += int(row["output_tokens"])
        result[key][f"{condition}_ms"] += float(row["median_ms"])
    return result


def make_benchmark(path: Path, font_dir: Path, rows):
    project = aggregate(rows, "project_id")
    display = {"cbm-c": "codebase-memory-mcp", "aikb-web": "aikb-web", "aikb-java": "aikb-java"}
    order = ["cbm-c", "aikb-web", "aikb-java"]
    baseline_total = sum(v["baseline_tokens"] for v in project.values())
    mcp_total = sum(v["mcp_tokens"] for v in project.values())
    baseline_ms = sum(v["baseline_ms"] for v in project.values())
    mcp_ms = sum(v["mcp_ms"] for v in project.values())

    w, h = 1600, 930
    image = Image.new("RGB", (w, h), WHITE)
    draw = ImageDraw.Draw(image)
    title_font = font(font_dir, 32, True)
    metric_font = font(font_dir, 42, True)
    body_font = font(font_dir, 22)
    small_font = font(font_dir, 19)
    draw.text((60, 40), "本机 3 项目 × 3 任务 A/B 实测", font=title_font, fill=NAVY)

    cards = [
        (60, 105, 500, 230, LIGHT_BLUE, f"{baseline_total:,} → {mcp_total:,}", "工具返回 Token", "减少 90.7%"),
        (580, 105, 1020, 230, LIGHT_TEAL, f"{baseline_ms:,.0f} → {mcp_ms:,.0f} ms", "直接 CLI 中位耗时合计", "减少 42.8%"),
        (1100, 105, 1540, 230, LIGHT_GOLD, "6 / 9 → 0 / 9", "达到 50,000 字符上限", "图谱结果无截断"),
    ]
    for x1, y1, x2, y2, fill, value, label, note in cards:
        rounded_box(draw, (x1, y1, x2, y2), fill, BORDER, radius=22, width=2)
        draw_center(draw, ((x1 + x2) / 2, y1 + 42), value, metric_font, NAVY)
        draw_center(draw, ((x1 + x2) / 2, y1 + 83), label, body_font, INK)
        draw_center(draw, ((x1 + x2) / 2, y1 + 108), note, small_font, MUTED)

    panels = [(60, 285, 770, 855, "返回 Token", "tokens"), (830, 285, 1540, 855, "查询耗时（ms）", "ms")]
    for x1, y1, x2, y2, label, metric in panels:
        rounded_box(draw, (x1, y1, x2, y2), WHITE, BORDER, radius=22, width=2)
        draw.text((x1 + 28, y1 + 22), label, font=title_font, fill=NAVY)
        draw.rectangle((x2 - 260, y1 + 28, x2 - 238, y1 + 50), fill=MUTED)
        draw.text((x2 - 228, y1 + 23), "纯文本", font=small_font, fill=INK)
        draw.rectangle((x2 - 135, y1 + 28, x2 - 113, y1 + 50), fill=TEAL)
        draw.text((x2 - 103, y1 + 23), "图谱", font=small_font, fill=INK)
        max_value = max(
            project[key][f"baseline_{metric}"] for key in order
        )
        chart_left = x1 + 205
        chart_right = x2 - 80
        row_y = y1 + 135
        for key in order:
            draw.text((x1 + 28, row_y + 30), display[key], font=small_font, fill=INK)
            for offset, condition, color in ((0, "baseline", MUTED), (38, "mcp", TEAL)):
                value = project[key][f"{condition}_{metric}"]
                bar_w = max(4, int((chart_right - chart_left) * value / max_value))
                y = row_y + offset
                draw.rounded_rectangle((chart_left, y, chart_left + bar_w, y + 25), radius=8, fill=color)
                value_label = f"{int(round(value)):,}" if metric == "ms" else f"{int(value):,}"
                label_x = min(chart_left + bar_w + 10, chart_right - 5)
                anchor = "la" if label_x < chart_right - 50 else "ra"
                draw.text((label_x, y + 12), value_label, font=small_font, fill=INK, anchor=anchor)
            row_y += 140
    draw.text((60, 887), "口径：o200k_base 精确计数；每项运行 3 次取中位数；索引构建成本未计入。", font=small_font, fill=MUTED)
    image.save(path, quality=95)


def add_picture(doc, path: Path, *, width=6.35, alt: str, caption: str):
    cp = doc.add_paragraph(style="Caption")
    cp.add_run(caption)
    cp.paragraph_format.keep_with_next = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    set_image_alt(shape, alt)
    return shape


def build(args):
    skill_scripts = Path(args.skill_root) / "scripts"
    sys.path.insert(0, str(skill_scripts))
    from table_geometry import apply_table_geometry

    output = Path(args.output).resolve()
    output.parent.mkdir(parents=True, exist_ok=True)
    work_dir = Path(args.work_dir).resolve()
    work_dir.mkdir(parents=True, exist_ok=True)
    font_dir = Path("C:/Windows/Fonts")
    source_dir = Path(__file__).resolve().parent
    core_rows = read_core_results(source_dir / "local-mcp-ab-core-results-2026-08-03.csv")
    task_data = aggregate(core_rows, "task")

    cover_path = work_dir / "competition-cover-graph.png"
    architecture_path = work_dir / "competition-architecture.png"
    benchmark_path = work_dir / "competition-benchmark.png"
    make_cover_graph(cover_path, font_dir)
    make_architecture(architecture_path, font_dir)
    make_benchmark(benchmark_path, font_dir, core_rows)

    doc = Document()
    set_document_geometry(doc)
    set_styles(doc)
    set_header_footer(doc)
    doc.core_properties.title = "Codebase Memory MCP 参赛项目汇报书"
    doc.core_properties.subject = "面向 AI 编码 Agent 的本地代码知识图谱"
    doc.core_properties.author = "Codebase Memory MCP 项目组"
    doc.core_properties.keywords = "MCP, 代码知识图谱, AI 编码 Agent, 静态分析, Token 优化"

    # Cover
    add_paragraph(doc, "技术创新与应用价值", style="Cover Kicker", before=12, after=18)
    title = add_paragraph(doc, "面向 AI 编码 Agent 的\n本地代码知识图谱", style="Cover Title", after=8, line=1.05)
    title.paragraph_format.keep_with_next = True
    add_paragraph(doc, "从逐文件搜索到可复用、可追踪、可审计的结构化代码记忆", style="Cover Subtitle", after=18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    shape = p.add_run().add_picture(str(cover_path), width=Inches(6.3))
    set_image_alt(shape, "代码实体、调用关系、服务接口、变更影响与测试链路组成的代码知识图谱")
    add_callout(
        doc,
        "将 AI 对代码库的一次性逐文件探索，升级为可复用的结构化理解，从而减少上下文成本、提高关系理解能力，并为代码变更建立可审计的影响证据。",
        label="价值主张",
        fill=LIGHT_BLUE,
        accent=BLUE,
        geometry=apply_table_geometry,
    )
    metadata = doc.add_table(rows=2, cols=2)
    meta_values = [
        ("项目定位", "基于 MIT 开源上游的 AI 编码 Agent 结构化代码记忆层"),
        ("技术路线", "Tree-sitter + Hybrid LSP + 知识图谱 + MCP"),
        ("本机实测", "3 个项目 / 9 个固定任务"),
        ("文档版本", "参赛版 V1.0 / 2026-08-05"),
    ]
    for idx, (label, value) in enumerate(meta_values):
        cell = metadata.cell(idx // 2, idx % 2)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(1)
        set_run_font(p.add_run(label + "\n"), size=8.5, bold=True, color=GOLD)
        set_run_font(p.add_run(value), size=9.5, color=INK)
    remove_table_borders(metadata)
    apply_table_geometry(metadata, [4680, 4680])

    # Page 2
    page_break(doc)
    add_section_title(doc, "评审速览", "项目摘要", "它不是代码搜索器，而是 AI Agent 与代码库之间的确定性结构分析后端。")
    add_metric_strip(
        doc,
        [
            ("90.7%", "返回 Token 减少", "109,334 → 10,194", LIGHT_BLUE),
            ("42.8%", "核心查询耗时减少", "2,944 → 1,685 ms", LIGHT_TEAL),
            ("3 × 3", "本机验证矩阵", "3 项目 / 3 类任务", LIGHT_GOLD),
        ],
        geometry=apply_table_geometry,
    )
    add_callout(
        doc,
        "Codebase Memory MCP 是基于 MIT 开源上游的当前 fork：将源码解析为持久化知识图谱，通过 MCP 向编码 Agent 提供搜索、调用链、影响分析、架构与跨服务证据；服务本身不内置 LLM。",
        label="项目定义",
        fill=LIGHT,
        accent=NAVY,
        geometry=apply_table_geometry,
    )
    doc.add_heading("评委最应关注的三件事", level=2)
    add_list(
        doc,
        [
            ("创新性：", "以 Tree-sitter 负责广覆盖语法提取，以 Hybrid LSP 修正跨文件类型关系，再以属性图进行确定性关系推理。"),
            ("实证性：", "9 组本机 A/B 使用 o200k_base 精确计数；文本基线已统一截断，Token 降幅仍达到 90.7%。"),
            ("可落地性：", "本地静态二进制、MCP/CLI/HTTP 多入口、SQLite 持久化和覆盖度元数据，使能力可嵌入个人开发、团队 CI 与企业知识治理。"),
        ],
    )
    doc.add_heading("预期价值", level=2)
    add_data_table(
        doc,
        ["对象", "当前成本", "项目带来的改变"],
        [
            ["开发者", "重复搜索、逐文件阅读", "用符号、路径和关系证据直接回答问题"],
            ["团队", "评审依赖经验、知识易流失", "沉淀可复用图谱、影响报告和架构快照"],
            ["企业", "多仓库依赖不可见、治理困难", "形成跨服务技术地图与可信代码知识门户"],
        ],
        [1500, 3000, 4860],
        geometry=apply_table_geometry,
        font_size=9.5,
    )

    # Page 3
    page_break(doc)
    add_section_title(doc, "1", "问题与机会", "AI 编码 Agent 的瓶颈，正在从“模型会不会写代码”转向“模型能否获得正确、足够且不过量的代码上下文”。")
    doc.add_heading("传统逐文件探索的五类成本", level=2)
    add_list(
        doc,
        [
            ("上下文成本高：", "目录、搜索结果和整文件内容被反复送入模型，Token 消耗随仓库规模快速增长。"),
            ("关系理解弱：", "文本命中只能说明字符串出现，不能可靠回答“谁调用它”“修改会影响谁”。"),
            ("跨文件语义不足：", "导入、继承、泛型、路径别名、方法分派与框架关系难以通过正则还原。"),
            ("理解结果不可复用：", "新会话、新成员和新 Agent 往往重复完成同一轮探索。"),
            ("可信边界缺失：", "结果为空时，无法区分“确实不存在”与“未索引、未解析或已过期”。"),
        ],
        kind="decimal",
    )
    doc.add_heading("从文件集合到知识资产", level=2)
    add_data_table(
        doc,
        ["维度", "传统文本探索", "知识图谱方式"],
        [
            ["检索对象", "字符串与文件", "函数、类、路由、配置和关系"],
            ["回答形式", "命中位置，需要继续判断", "限定名、行号、路径、调用边与覆盖状态"],
            ["上下文", "大段原文与重复命中", "与问题相关的最小结构化证据集"],
            ["复用能力", "会话级、一次性", "持久化、增量更新、团队共享"],
            ["风险判断", "依赖人工拼接", "可追踪直接/间接影响与测试链路"],
        ],
        [1600, 3600, 4160],
        geometry=apply_table_geometry,
        font_size=9.3,
        header_fill=LIGHT_BLUE,
    )
    add_callout(
        doc,
        "机会不在于用更复杂的方式完成一次搜索，而在于把代码结构预计算为可持续复用的“上下文基础设施”。",
        label="关键判断",
        fill=LIGHT_GOLD,
        accent=GOLD,
        geometry=apply_table_geometry,
    )

    # Page 4
    page_break(doc)
    add_section_title(doc, "2", "解决方案与总体架构", "通过“语法提取 + 语义修正 + 图关系构建 + 标准工具协议”，把代码资产转换为 Agent 可直接消费的证据。")
    add_picture(
        doc,
        architecture_path,
        alt="源代码经过文件发现、Tree-sitter、Hybrid LSP 与关系构建形成持久化知识图谱，再通过 MCP、CLI、HTTP 和图形界面提供查询能力",
        caption="图 1  Codebase Memory MCP 端到端技术架构",
    )
    doc.add_heading("四项架构原则", level=2)
    add_list(
        doc,
        [
            ("结构优先：", "先识别定义、调用、导入、继承和路由，再把相关源码作为证据回读。"),
            ("语法与语义分层：", "Tree-sitter 保证广覆盖，Hybrid LSP 提高跨文件与类型相关关系质量。"),
            ("本地与模型解耦：", "服务不内置 LLM，不需要额外 API Key，查询结果可被不同 MCP 客户端复用。"),
            ("可信边界显式化：", "索引模式、解析缺口、文件新鲜度和分页状态与答案一并返回。"),
        ],
    )

    # Page 5
    page_break(doc)
    add_section_title(doc, "3", "核心创新与技术壁垒", "项目的壁垒不是单点算法，而是多层代码理解能力在一个本地、可查询、可审计系统中的组合。")
    innovations = [
        ("3.1  双层解析架构", "Tree-sitter 从广泛语言中稳定提取 AST 结构；Hybrid LSP 进一步推导导入、泛型、继承、返回类型和方法分派，修正 CALLS、USAGE 等关系。"),
        ("3.2  确定性图关系 + 语义检索", "以限定名、调用图和类型边保证可解释关系；同时结合 BM25、驼峰分词、语义向量和相似性信号，弥合自然语言与代码命名差异。"),
        ("3.3  面向变更的影响传播", "从目标符号或 Git diff 反向遍历直接/间接依赖，连接入口、测试和跨服务边，为评审与回归提供证据。"),
        ("3.4  MCP 关注点分离", "Agent 负责理解自然语言意图，图服务负责确定性执行；同一个知识层可服务多种模型、IDE 和自动化流程。"),
        ("3.5  覆盖度与新鲜度模型", "系统不仅返回命中，还说明未索引文件、部分解析区间、元数据变化和回退源码建议，降低“未知被误报为不存在”的风险。"),
    ]
    for heading, body in innovations:
        doc.add_heading(heading, level=2)
        p = add_paragraph(doc, body, after=6, line=1.23)
        p.paragraph_format.keep_together = True
    add_callout(
        doc,
        "创新组合：广覆盖 AST + 轻量类型语义 + 持久化属性图 + 标准 MCP 协议 + 覆盖度证据。",
        label="技术壁垒",
        fill=LIGHT_TEAL,
        accent=TEAL,
        geometry=apply_table_geometry,
    )

    # Page 6
    page_break(doc)
    add_section_title(doc, "4", "本机 A/B 实测", "用可复现的真实仓库数据证明价值，并将项目文档口径与本次独立复测严格区分。")
    add_picture(
        doc,
        benchmark_path,
        alt="三个本机项目中纯文本与图谱核心查询的 Token 和查询耗时对比",
        caption="图 2  3 个本机项目、9 个固定任务的 A/B 结果",
    )
    add_data_table(
        doc,
        ["项目", "纯文本 Token", "图谱 Token", "降幅", "文本 ms", "图谱 ms"],
        [
            ["codebase-memory-mcp", "7,618", "3,088", "59.5%", "386.15", "372.75"],
            ["aikb-web", "49,190", "1,840", "96.3%", "1,741.87", "588.15"],
            ["aikb-java", "52,526", "5,266", "90.0%", "816.32", "723.78"],
        ],
        [2400, 1440, 1350, 1100, 1535, 1535],
        geometry=apply_table_geometry,
        font_size=9.1,
        numeric_cols=(1, 2, 3, 4, 5),
        header_fill=LIGHT_BLUE,
    )
    p = add_paragraph(
        doc,
        "数据口径：本机 2026-08-03 运行；每项 3 次取中位数；o200k_base 精确计数；两边统一 50,000 字符捕获上限；复用已有索引。",
        style="Source Note",
        after=2,
        line=1.1,
    )

    # Page 7
    page_break(doc)
    add_section_title(doc, "5", "结果解读与可信边界", "价值必须同时看上下文效率、核心查询性能、答案语义和当前客户端体验。")
    add_data_table(
        doc,
        ["任务", "纯文本 Token", "图谱 Token", "降幅", "文本 ms", "图谱 ms"],
        [
            ["定义定位", f"{task_data['definition']['baseline_tokens']:,}", f"{task_data['definition']['mcp_tokens']:,}", "99.1%", "974.41", "335.19"],
            ["两跳关系追踪", f"{task_data['relations']['baseline_tokens']:,}", f"{task_data['relations']['mcp_tokens']:,}", "81.2%", "1,026.06", "310.54"],
            ["架构摸底", f"{task_data['architecture']['baseline_tokens']:,}", f"{task_data['architecture']['mcp_tokens']:,}", "91.6%", "943.87", "1,038.95"],
        ],
        [2100, 1460, 1360, 1100, 1670, 1670],
        geometry=apply_table_geometry,
        font_size=9.1,
        numeric_cols=(1, 2, 3, 4, 5),
        header_fill=LIGHT_TEAL,
    )
    doc.add_heading("如何正确表述效率", level=2)
    add_list(
        doc,
        [
            ("核心查询：", "直接 CLI 的中位耗时合计减少 42.8%，说明图查询引擎具备实用性能。"),
            ("上下文：", "返回 Token 减少 90.7%；未截断的文本原始输出超过 102 万 Token，因此主结果已采用更保守基线。"),
            ("答案语义：", "图谱返回限定名和已解析关系；纯文本关系任务不是同质量输出，本次未用 LLM Judge，因此不夸大答案质量。"),
            ("集成体验：", "Codex 端到端补测仍观察到明显额外等待，应通过连接复用和分段遥测定位客户端、协议与宿主开销。"),
        ],
    )
    add_callout(
        doc,
        "比赛汇报建议采用“90.7% Token 减少 + 42.8% 核心查询耗时减少”作为主证据，同时主动说明索引成本、静态分析边界和 Codex 集成链路仍需优化。",
        label="可信口径",
        fill=LIGHT_GOLD,
        accent=GOLD,
        geometry=apply_table_geometry,
    )

    # Page 8
    page_break(doc)
    add_section_title(doc, "6", "典型应用与演示闭环", "项目的价值从“回答一个代码问题”延伸到评审、排障、迁移和组织级知识治理。")
    applications = [
        ("新成员上手", "快速获得语言分布、模块边界、入口、路由、热点和核心调用链。"),
        ("代码修改与评审", "在修改公共函数或组件前生成调用方、受影响文件、入口、测试和跨服务影响。"),
        ("故障排查", "从接口或异常方法追踪到控制器、业务逻辑、存储、配置、下游服务和事件 Channel。"),
        ("技术债治理", "识别高复杂度函数、近似重复、死代码候选、高扇入节点和边界异常。"),
        ("遗留系统现代化", "为 Vue 2 / Webpack 升级、框架迁移和模块拆分形成可解释的任务批次。"),
        ("AI 编码 Agent 增强", "为重构、测试生成和缺陷修复提供结构化最小证据集。"),
    ]
    add_data_table(
        doc,
        ["应用场景", "可交付结果"],
        applications,
        [2200, 7160],
        geometry=apply_table_geometry,
        font_size=9.6,
        header_fill=LIGHT,
    )
    doc.add_heading("建议现场演示：8 分钟价值链", level=2)
    add_list(
        doc,
        [
            "索引一个真实 Vue 2 业务仓库，展示覆盖状态与新鲜度。",
            "用 get_architecture 展示模块、入口、路由与热点。",
            "用 search_graph 定位一个公共组件或业务函数。",
            "用 trace_path 展示上下游调用链。",
            "用 explain_impact 输出受影响文件、入口和测试。",
            "用 get_code_snippet 回读源码，并在 3D UI 展示跨模块关系。",
        ],
        kind="decimal",
    )

    # Page 9
    page_break(doc)
    add_section_title(doc, "7", "竞争差异与可持续优势", "项目不与文本搜索、IDE LSP 或向量 RAG 二选一，而是把它们的优势整合为结构化代码理解基础设施。")
    add_data_table(
        doc,
        ["方案", "优势", "主要局限", "本项目的补充价值"],
        [
            ["grep / ripgrep", "快速、精确字面匹配", "不理解类型、调用和影响", "将文本命中提升为符号与关系证据"],
            ["IDE LSP", "单语言语义准确", "多语言部署复杂，跨仓持久化弱", "统一图模型、持久化关系和 MCP 接口"],
            ["向量 RAG", "擅长语义召回", "确定性关系与调用链较弱", "语义检索与图关系联合"],
            ["内置 LLM 平台", "自然语言体验完整", "额外模型、成本与数据边界", "复用已有 Agent，本地、模型无关"],
            ["人工架构文档", "表达清晰、便于沟通", "容易过期，维护成本高", "从当前代码持续生成，并以 ADR 补充意图"],
        ],
        [1600, 2050, 2600, 3110],
        geometry=apply_table_geometry,
        font_size=8.8,
        header_fill=LIGHT_BLUE,
    )
    doc.add_heading("开源底座与本分支贡献", level=2)
    add_data_table(
        doc,
        ["归属层", "边界与可对外表述"],
        [
            ["开源上游（MIT）", "DeusData/codebase-memory-mcp 提供 Tree-sitter 多语言解析、知识图谱与 SQLite 持久化、结构化查询，以及 MCP/CLI/HTTP 基础能力；本项目遵循上游许可证并保留归属。"],
            ["当前 fork 新增/强化", "ycsx/codebase-memory-mcp 在上游基础上开发或强化 Vue 2 单文件组件与 Webpack 别名分析、本地化图谱控制台、Windows Desktop 控制与更新流程、跨平台签名发布和 CI/诊断工程。"],
            ["本次参赛交付", "当前团队完成 3 个项目 × 3 类任务的 Windows 本机 A/B 设计与复测，形成 Token/效率报告、Vue 2 / Webpack 演示闭环和本参赛汇报材料。"],
        ],
        [2050, 7310],
        geometry=apply_table_geometry,
        font_size=8.7,
        header_fill=LIGHT_TEAL,
    )
    advantages_heading = doc.add_heading("三类可持续优势", level=2)
    advantages_heading.paragraph_format.space_before = Pt(8)
    advantages_heading.paragraph_format.space_after = Pt(4)
    advantage_paragraphs = add_list(
        doc,
        [
            ("数据资产效应：", "索引可持久化、增量更新和团队共享，查询越多，首次构建成本越被摊薄。"),
            ("生态接口效应：", "MCP 让同一结构分析后端服务不同模型、IDE、Agent 和自动化流程。"),
            ("领域规则效应：", "通过 Hybrid LSP、框架解析和跨服务规则持续提高特定技术栈的关系质量。"),
        ],
    )
    for paragraph in advantage_paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.12

    # Page 10
    add_section_title(
        doc,
        "8",
        "当前分支特色与落地切口",
        "以 Vue 2 / Webpack 遗留前端治理为差异化入口，把通用图谱能力转化为具体业务价值。",
        page_break_before=True,
    )
    add_metric_strip(
        doc,
        [
            ("14,674", "本仓图节点", "自举分析结果", LIGHT_BLUE),
            ("84,329", "本仓图关系", "含调用、使用与相似性", LIGHT_TEAL),
            ("586", "索引文件", "多语言工程资产", LIGHT_GOLD),
        ],
        geometry=apply_table_geometry,
    )
    doc.add_heading("Vue 单文件组件深度解析", level=2)
    add_paragraph(doc, "系统提取 .vue 文件中的 script / script setup 内容，继续生成定义、导入、函数调用和 Channel 关系，并映射回原始 Vue 文件行号。", after=6, line=1.23)
    doc.add_heading("Vue CLI 与 Webpack 路径别名", level=2)
    add_paragraph(doc, "静态读取 tsconfig/jsconfig 的 paths 与 baseUrl、Vue CLI 默认 @ 别名、vue.config 和 webpack.config 中的静态 alias；分析过程中不执行配置代码。", after=6, line=1.23)
    doc.add_heading("可形成的专项产品", level=2)
    add_list(
        doc,
        [
            "公共组件修改影响与依赖盘点。",
            "失效别名、动态加载和路由关系扫描。",
            "Vue 2 到 Vue 3、Webpack 到 Vite 的迁移评估与批次拆分。",
            "前端组件到后端 API 的跨层关系追踪。",
        ],
    )
    add_callout(
        doc,
        "通用能力负责“看懂代码关系”，Vue 2 / Webpack 专项负责“解决一个存量巨大、迁移迫切的真实市场问题”。",
        label="落地策略",
        fill=LIGHT_TEAL,
        accent=TEAL,
        geometry=apply_table_geometry,
    )

    # Page 11
    page_break(doc)
    add_section_title(doc, "9", "产品化路线与价值扩张", "以查询效率为起点，向团队研发治理和企业技术资产地图逐级扩张。")
    add_data_table(
        doc,
        ["阶段", "重点建设", "交付物", "目标价值"],
        [
            ["P0  可用与可信", "端到端遥测、连接复用、分页协议、口径校准", "自动 A/B、性能分解、覆盖门禁", "让 ROI 可测、结果可审计"],
            ["P1  团队工作流", "PR/CI 影响分析、测试推荐、Vue 专项", "风险报告、回归清单、迁移评估", "进入评审和交付主流程"],
            ["P2  企业知识层", "运行时 Trace、时序图谱、Owner/ADR/事故连接", "架构漂移、技术地图、知识门户", "形成组织级研发资产"],
        ],
        [1450, 3050, 2700, 2160],
        geometry=apply_table_geometry,
        font_size=9.1,
        header_fill=LIGHT_BLUE,
    )
    doc.add_heading("价值扩张路径", level=2)
    add_list(
        doc,
        [
            ("个人开发者：", "减少 Agent 的无效检索与上下文消耗，提高重构和排障效率。"),
            ("研发团队：", "在 PR/CI 中自动产生影响面、测试建议和架构风险。"),
            ("企业平台：", "连接多仓库、服务目录、Owner、ADR 与事故，构建技术资产地图。"),
        ],
    )
    doc.add_heading("可持续运营方式", level=2)
    add_paragraph(doc, "基础能力保持本地、开放和模型无关；团队版围绕共享图谱、权限、审计与 CI 集成；企业版围绕多仓库治理、运行时校准和组织知识连接形成增值。", after=8, line=1.25)
    add_callout(
        doc,
        "产品化不应追求“所有问题都走图谱”，而应建设自适应路由：字面量与配置检索走文本，关系、影响、架构和跨服务问题走图谱。",
        label="产品原则",
        fill=LIGHT_GOLD,
        accent=GOLD,
        geometry=apply_table_geometry,
    )

    # Page 12
    page_break(doc)
    add_section_title(doc, "10", "风险边界与治理方案", "主动披露边界不是削弱项目，而是让静态分析结论能够进入真实工程决策。")
    add_data_table(
        doc,
        ["风险", "影响", "治理方案"],
        [
            ["静态分析盲区", "反射、动态注册和运行时代码可能漏解析", "接入 OpenTelemetry/测试 Trace；关键结论回读源码"],
            ["索引新鲜度", "图谱可能落后于工作区", "watcher、元数据比对、覆盖门禁和自动增量更新"],
            ["大结果截断", "局部结果可能被误报为全集", "统一 has_more、总数、游标和分页审计"],
            ["客户端链路延迟", "核心查询快但真实等待仍高", "拆分客户端、传输、进程、查询和序列化遥测"],
            ["基准外推", "单机小样本不能代表所有仓库", "扩大分层问题集，记录完整会话 Token 和质量评分"],
            ["能力归属", "上游能力与当前 fork 贡献可能混淆", "在文档、演示和发布说明中明确区分"],
        ],
        [2100, 2950, 4310],
        geometry=apply_table_geometry,
        font_size=9.0,
        header_fill=LIGHT_GOLD,
    )
    doc.add_heading("可信分析四层证据", level=2)
    add_list(
        doc,
        [
            ("结论：", "回答定义、路径、调用和影响。"),
            ("定位：", "返回 qualified name、文件与行号。"),
            ("覆盖：", "返回未索引、部分解析和新鲜度状态。"),
            ("回退：", "对高风险或覆盖不足范围建议源码核验。"),
        ],
        kind="decimal",
    )
    add_callout(
        doc,
        "项目的可信性来自“答案 + 证据 + 覆盖边界 + 回退动作”，而不是承诺静态图绝对完整。",
        label="可信设计",
        fill=LIGHT_BLUE,
        accent=BLUE,
        geometry=apply_table_geometry,
    )

    # Page 13
    page_break(doc)
    add_section_title(doc, "11", "结论：为什么值得参赛", "Codebase Memory MCP 将编译原理、静态分析、图数据库和 MCP 生态组合成一个可验证、可交付、可扩展的 AI 编码基础设施。")
    add_data_table(
        doc,
        ["评审维度", "项目表现"],
        [
            ["原创性", "提出 AI Agent 与代码库之间的“结构化代码记忆层”，以语法、语义和图关系联合供给上下文。"],
            ["技术深度", "覆盖 AST、Hybrid LSP、图遍历、影响传播、语义检索、增量索引与覆盖度模型。"],
            ["量化价值", "本机 9 项 A/B：返回 Token 减少 90.7%，核心查询中位耗时合计减少 42.8%。"],
            ["工程成熟度", "本地二进制、SQLite 持久化、MCP/CLI/HTTP、跨平台发布设计、3D UI 与诊断能力。"],
            ["应用前景", "可从个人 Agent 增强扩展到 PR 风险门禁、遗留系统治理与组织级技术地图。"],
            ["可信边界", "显式披露覆盖、新鲜度、分页和静态分析限制，适合进入工程决策流程。"],
        ],
        [1900, 7460],
        geometry=apply_table_geometry,
        font_size=9.5,
        header_fill=LIGHT_BLUE,
    )
    quote = add_paragraph(
        doc,
        "让 AI 不再反复“阅读代码库”，而是持续“理解代码关系”。",
        style="Closing Quote",
        before=16,
        after=14,
        line=1.2,
    )
    quote.paragraph_format.keep_together = True
    add_callout(
        doc,
        "短期，它降低 Agent 的 Token 和探索成本；中期，它进入评审、测试和迁移流程；长期，它将代码、架构与组织知识连接为可持续演化的研发数字资产。",
        label="参赛价值",
        fill=LIGHT_TEAL,
        accent=TEAL,
        geometry=apply_table_geometry,
    )
    doc.add_heading("数据来源与参考", level=2)
    source_paragraphs = add_list(
        doc,
        [
            "复测数据与方法：docs/reports/local-mcp-ab-core-results-2026-08-03.csv；docs/reports/LOCAL_MCP_AB_BENCHMARK_2026-08-03.md。",
            "项目口径：README.md；docs/PROJECT_REPORT.md；Codebase-Memory: Tree-Sitter-Based Knowledge Graphs for LLM Code Exploration via MCP，arXiv:2603.27277。",
        ],
    )
    for paragraph in source_paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            set_run_font(run, size=8.5, color=MUTED)
    add_paragraph(
        doc,
        "说明：项目 README 中的跨仓库论文数据、语言数量和 Apple M3 Pro 性能属于项目文档口径；本汇报的 90.7% 与 42.8% 为当前 Windows 本机独立复测结果。",
        style="Source Note",
        after=0,
        line=1.1,
    )

    doc.save(output)
    print(output)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", required=True)
    parser.add_argument("--work-dir", required=True)
    parser.add_argument("--skill-root", required=True)
    args = parser.parse_args()
    build(args)


if __name__ == "__main__":
    main()
