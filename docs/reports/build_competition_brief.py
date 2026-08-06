from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT_CJK = "Microsoft YaHei"
FONT_LATIN = "Calibri"
NAVY = "102A43"
BLUE = "2E74B5"
TEAL = "167C8C"
MUTED = "607080"
LIGHT_BLUE = "E8EEF5"
LIGHT_TEAL = "E7F3F2"
LIGHT_GOLD = "FFF4D9"
LIGHT_GRAY = "F4F6F8"
GRID = "D4DEE8"
WHITE = "FFFFFF"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, *, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CJK)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), FONT_LATIN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def set_style_font(style, *, size, color=NAVY, bold=False, line=1.1, after=0):
    style.font.name = FONT_LATIN
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_LATIN)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CJK)
    style._element.get_or_add_rPr().rFonts.set(qn("w:cs"), FONT_LATIN)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.line_spacing = line
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(after)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, color=GRID, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_borders(cell)


def set_keep(paragraph, *, next=False, lines=True):
    paragraph.paragraph_format.keep_with_next = next
    paragraph.paragraph_format.keep_together = lines


def add_text(paragraph, text, *, size=10.5, color=NAVY, bold=False, italic=False):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return run


def add_paragraph(doc, text="", *, size=10.5, color=NAVY, bold=False, after=4, line=1.1, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    add_text(p, text, size=size, color=color, bold=bold)
    return p


def add_heading(doc, text, *, level=1, after=4, before=6):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    set_keep(p, next=True)
    add_text(p, text, size=16 if level == 1 else 12.2, color=BLUE, bold=True)
    return p


def add_metric_strip(doc):
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    set_table_geometry(table, [3120, 3120, 3120])
    data = [
        ("90.7%", "返回上下文 Token 减少", LIGHT_BLUE),
        ("42.8%", "直接 CLI 中位耗时减少", LIGHT_TEAL),
        ("0 / 9", "图谱结果触达 50,000 字符上限", LIGHT_GOLD),
    ]
    for cell, (value, label, fill) in zip(table.rows[0].cells, data):
        set_cell_shading(cell, fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        add_text(p, value, size=18, color=NAVY, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1.0
        add_text(p2, label, size=8.5, color=MUTED)
    return table


def add_flow(doc):
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    set_table_geometry(table, [2340, 2340, 2340, 2340])
    steps = [
        ("代码仓库", "输入"),
        ("Tree-sitter + Hybrid LSP", "解析"),
        ("知识图谱", "关系"),
        ("MCP / CLI / Agent", "输出证据"),
    ]
    for cell, (title, desc) in zip(table.rows[0].cells, steps):
        set_cell_shading(cell, LIGHT_GRAY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        add_text(p, title, size=9.3, color=NAVY, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1.0
        add_text(p2, desc, size=8.1, color=MUTED)
    return table


def add_highlights(doc):
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    set_table_geometry(table, [3120, 3120, 3120])
    blocks = [
        ("个人开发效率", "减少手工搜索与上下文拼接，让排障、重构和新成员上手更快。"),
        ("团队交付质量", "把影响面、测试建议和架构风险接入 PR/CI，减少遗漏和返工。"),
        ("遗留系统治理", "为 Vue 2 / Webpack 提供别名、依赖和迁移评估，缩短现代化改造路径。"),
    ]
    for cell, (title, body) in zip(table.rows[0].cells, blocks):
        set_cell_shading(cell, WHITE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        add_text(p, title, size=10.4, color=TEAL, bold=True)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1.05
        add_text(p2, body, size=8.7, color=NAVY)
    return table


def add_attribution(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    add_text(p, "成果边界", size=10.2, color=BLUE, bold=True)
    table = doc.add_table(rows=3, cols=2)
    table.autofit = False
    set_table_geometry(table, [1850, 7510])
    rows = [
        ("开源上游（MIT）", "DeusData/codebase-memory-mcp：Tree-sitter 多语言解析、图谱与 SQLite 持久化、结构化查询和 MCP/CLI/HTTP 基础能力。"),
        ("当前 fork 新增/强化", "Vue 2 SFC 与 Webpack 别名分析、本地化图谱控制台、Windows Desktop 控制/更新、跨平台签名发布与 CI 诊断。"),
        ("本次参赛交付", "Windows 本机 3 项目 × 3 任务 A/B 测试、Token/效率对比、Vue 2 / Webpack 演示闭环及本汇报材料。"),
    ]
    fills = [LIGHT_GRAY, LIGHT_TEAL, LIGHT_GOLD]
    for row, (label, body), fill in zip(table.rows, rows, fills):
        set_cell_shading(row.cells[0], fill)
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for idx, cell in enumerate(row.cells):
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            add_text(p, label if idx == 0 else body, size=7.7 if idx == 0 else 7.5, color=NAVY, bold=idx == 0)
    return table


def set_alt_text(inline_shape, description):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description)


def configure_document(doc):
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    set_style_font(normal, size=10, color=NAVY, line=1.1, after=4)
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 6, 4),
        ("Heading 2", 12.2, BLUE, 5, 3),
        ("Heading 3", 11, TEAL, 4, 2),
    ]:
        style = doc.styles[name]
        set_style_font(style, size=size, color=color, bold=True, line=1.0, after=after)
        style.paragraph_format.space_before = Pt(before)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hp.paragraph_format.line_spacing = 1.0
    add_text(hp, "CODEBASE MEMORY MCP", size=8.2, color=MUTED, bold=True)
    add_text(hp, "    评审速览版", size=8.2, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    fp.paragraph_format.line_spacing = 1.0
    add_text(fp, "Windows 本机 A/B 复测｜o200k_base 精确计数｜完整版另附", size=7.4, color=MUTED)


def build(output: Path):
    doc = Document()
    configure_document(doc)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(3)
    kicker.paragraph_format.space_after = Pt(3)
    kicker.paragraph_format.line_spacing = 1.0
    add_text(kicker, "AI 编码基础设施 / COMPETITION BRIEF", size=8.2, color=TEAL, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(2)
    title.paragraph_format.line_spacing = 1.0
    set_keep(title, next=True)
    add_text(title, "Codebase Memory MCP", size=25, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(6)
    subtitle.paragraph_format.line_spacing = 1.05
    set_keep(subtitle, next=True)
    add_text(subtitle, "让 AI Agent 先理解代码关系，再回答问题", size=14.5, color=BLUE, bold=True)

    callout = doc.add_table(rows=1, cols=1)
    callout.autofit = False
    set_table_geometry(callout, [PAGE_WIDTH_DXA])
    cell = callout.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    add_text(p, "它是什么：", size=10.4, color=TEAL, bold=True)
    add_text(p, "一个运行在本机的 MCP 后端，把代码库解析成可查询的知识图谱，向 Agent 提供定义、调用链、架构和影响范围等结构化证据。", size=9.3, color=NAVY)

    add_paragraph(doc, "从“反复搜索整文件”变成“查询关系并回读证据”", size=9.4, color=MUTED, bold=True, after=3, line=1.0)
    add_flow(doc)
    add_metric_strip(doc)

    value_note = doc.add_paragraph()
    value_note.paragraph_format.space_before = Pt(3)
    value_note.paragraph_format.space_after = Pt(0)
    value_note.paragraph_format.line_spacing = 1.0
    add_text(value_note, "价值不止是省 Token：", size=8.7, color=TEAL, bold=True)
    add_text(value_note, "把一次性代码问答沉淀成团队可复用、可审计的技术关系资产。", size=8.7, color=NAVY)

    chart_path = Path(__file__).resolve().parents[2] / "build" / "competition-report" / "competition-benchmark.png"
    if chart_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        shape = p.add_run().add_picture(str(chart_path), width=Inches(6.35))
        set_alt_text(shape, "本机 3 项目乘 3 任务 A/B 对比：图谱减少返回 Token、查询耗时和字符上限触达")
    else:
        add_paragraph(doc, "A/B 图表未找到，请先生成 build/competition-report/competition-benchmark.png。", size=8.5, color=MUTED, after=2)

    add_heading(doc, "价值落点", level=2, before=2, after=2)
    add_highlights(doc)
    add_attribution(doc)
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(3)
    note.paragraph_format.space_after = Pt(0)
    note.paragraph_format.line_spacing = 1.0
    add_text(note, "归属说明：上游能力按开源许可标注；个人原创范围应以可核验的提交记录或 PR 记录为准。", size=7.1, color=MUTED, italic=True)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(output)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", required=True)
    args = parser.parse_args()
    build(Path(args.output))


if __name__ == "__main__":
    main()
